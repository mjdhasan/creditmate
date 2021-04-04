#### creditmate
# app_id: 17910ef6-8644-48b7-8dbb-a98e16298d76
# client secret id: 62ad4b96-1208-40a6-9395-a794730a81b1
# client secret value: mt14z3d-TsLyZe_uN2aQZ_jb1PKyMNVF~G

import onedrivesdk
from onedrivesdk.helpers import GetAuthCodeServer
import pickle


def navigate(client, item_id):
    items = client.item(id=item_id).children.get()
    return items

## CreditMate (majid@marketojo.com): https://portal.azure.com/#blade/Microsoft_AAD_RegisteredApps/ApplicationMenuBlade/Overview/appId/17910ef6-8644-48b7-8dbb-a98e16298d76/isMSAApp/
redirect_uri = 'https://127.0.0.1:5500/auth/redirect'
client_id = '17910ef6-8644-48b7-8dbb-a98e16298d76'
client_secret = 'mt14z3d-TsLyZe_uN2aQZ_jb1PKyMNVF~G'
tenant_id = 'b9fac4ac-dde5-4983-a2bc-ab447e7cfb6f'
## creditmate (majid.ubc@live.com): https://portal.azure.com/#blade/Microsoft_AAD_RegisteredApps/ApplicationMenuBlade/Overview/appId/17910ef6-8644-48b7-8dbb-a98e16298d76/isMSAApp/
# client_id = '1e0327a2-c117-45e0-a5d7-d87aceb563e2',
# client_secret = 'mqOeR7u6JYswffm.KQt-60UjD.C8R.El3g'
scopes=['wl.signin', 'wl.offline_access', 'onedrive.readwrite']
client = onedrivesdk.get_default_client(client_id=client_id, scopes=scopes)
# auth_url = client.auth_provider.get_auth_url(redirect_uri)
#this will block until we have the code
# code = GetAuthCodeServer.get_auth_code(auth_url, redirect_uri)  # M.R3_BAY.95cb5802-8033-a941-3b8d-1cece24690a4
code = 'M.R3_BAY.9b591cc9-ba5a-0b3e-0e9d-c7f990492e4a'
client.auth_provider.authenticate(code, redirect_uri, client_secret)

# with open('credentials/onedrive_client.pkl', 'wb') as output:
#     pickle.dump(client, output)

if __name__ == '__main__':

    l_drives = [drive for drive in client.drives.get()]
    drive = l_drives[0]
    dir(drive)
    drive.to_dict()
    drive.items

    client.auth_provider
    client.base_url
    dir(client.drive)
    item_id = "root"
    items = navigate(client, item_id)
    dir(items[4])
    [item.name for item in items]
    items[3].web_url
    items[4].name
    dir(items[4].file)
    items[4].file.to_dict()
    items[1].id
    dir(client.item(id=items[4].id))

    def parse_drive_tree(id, drive_dict, lvl):
        print(f'id: {id}; lvl: {lvl}')
        print(drive_dict)
        if f'l{lvl}' not in drive_dict.keys():
            drive_dict[f'l{lvl}'] = {}
        drive_dict[f'l{lvl}'][id] = [child for child in client.item(id=id).children.get()]
        for child in drive_dict[f'l{lvl}'][id]:
            parse_drive_tree(child.id, drive_dict, lvl + 1)
            # l_grandchild = [child for child in client.item(id=child.id).children.get()]
            # if len(l_grandchild) > 0:
            #     parse_drive_tree(id=, drive_dict, lvl + 1)

        return drive_dict


    drive_dict = parse_drive_tree(id='root', drive_dict={}, lvl=0)


    drive_dict = {}
    drive_dict['l0'] = [child for child in client.item(id='root').children.get()]
    drive_dict['l1'] = {}
    for child in drive_dict['l0']:
        drive_dict['l1'][child.id] = [child for child in client.item(id=child.id).children.get()]
    drive_dict['l2'] = {}
    for child in drive_dict['l1']:
        for grandchild in drive_dict['l1'][child]:
            drive_dict['l2'][grandchild.id] = [child for child in client.item(id=grandchild.id).children.get()]

    with open('data/onedrive_dict.pkl', 'wb') as output:
        pickle.dump(drive_dict, output, pickle.HIGHEST_PROTOCOL)

    # client.item(id=items[1].id).children.get()
    # dir(client.item(id=items[1].id).children)
    file_download = drive_dict['l1']['6CB02B0C51E92B1C!3443'][0]
    client.item(id=file_download.id).download(f'data/{file_download.name}')

    import docx
    doc = docx.Document(f'data/{file_download.name}')
    dir(doc)
    dir(doc.paragraphs[0])
    doc.paragraphs[0].text

    # # from some stackexchange answer i guess
    # import requests
    # # Build the POST parameters
    # params = {
    #           'grant_type': 'refresh_token',
    #           'client_id': '17910ef6-8644-48b7-8dbb-a98e16298d76',
    #           'refresh_token': 'M.R3_BAY.d3ec89a5-fa95-347f-9a7d-29a5c1b04541'
    #          }
    # response = requests.post('https://login.microsoftonline.com/common/oauth2/v2.0/token', data=params)
    # access_token = response.json()['access_token']
    # new_refresh_token = response.json()['refresh_token']
    # # ^ Save somewhere the new refresh token.
    # # I just overwrite the file with the new one.
    # # This new one will be used next time.
    # header = {'Authorization': 'Bearer ' + access_token}
    # # Download the file
    # response = requests.get('https://graph.microsoft.com/v1.0/me/drive/root:' +
    #                          PATH_TO_FILE + '/' + FILE_NAME + ':/content', headers=header)
    # # Save the file in the disk
    # with open(file_name, 'wb') as file:
    #     file.write(response.content)
    #
    #
    # import cloudsync
    # oauth_config = cloudsync.command.utils.generic_oauth_config('onedrive')
    # prov = cloudsync.create_provider('onedrive', oauth_config=oauth_config)
    # creds = prov.authenticate()
    #
    # # cdata
    # import cdata.onedrive as mod